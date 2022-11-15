#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
@author:    levondang
@contact:   levondang@163.com
@project:   beisong
@file:      to_docx.py
@time:      2022-11-13 14:09
@license:   Apache Licence
"""
from docx import Document
from docx.shared import Inches
import pandas
# https://blog.csdn.net/lly1122334/article/details/109669667

document = Document()
document.add_heading('北宋面试公众号推文收录', level=0)  # 插入标题

posts = pandas.read_csv("./beisong_pages.csv")
n = len(posts)
column_name = ["datetime", "url", "title", "miaoshu", "yaoqiu", "daan"]

for i in range(n):
    datetime, url, title, miaoshu, yaoqiu, daan = posts.loc[i, ["datetime"]].values[0], posts.loc[i, ["url"]].values[0], posts.loc[i, ["title"]].values[0], posts.loc[i, ["miaoshu"]].values[0], posts.loc[i, ["yaoqiu"]].values[0], posts.loc[i, ["daan"]].values[0]
    daan = daan.replace("\r", "")
    document.add_heading(f'{title}', level=1)  # 插入标题 1
    document.add_paragraph(f"datetime: {datetime}")
    document.add_paragraph(f"url: {url}")

    document.add_paragraph(f"{miaoshu}")
    document.add_paragraph(f"{yaoqiu}")
    document.add_paragraph(f"{daan}")

document.save(f"北宋面试公众号推文收录{n}篇.docx")
pass
